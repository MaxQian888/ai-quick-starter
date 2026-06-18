#!/usr/bin/env python3
"""把 Markdown 渲染成排版整洁的 .docx —— 简历 / 求职信 / 项目讲稿通用。

为什么存在这个脚本：Job Hunt Copilot 的交付环节需要 .docx，而旧版让模型去读
`/mnt/skills/public/docx/SKILL.md`，那个路径只在 Anthropic 托管环境存在，本地 /
Windows / Claude Code 下根本没有，交付会直接卡住。这个脚本自带 python-docx（缺了
自动装），在任何能跑 Python 的环境都能产出有效文件。

模型只需写它最擅长的 Markdown，再调用本脚本即可，无需手搓 docx API。

用法:
    python build_docx.py --in resume.md --out resume.docx [--style resume|letter|pitch]
    echo "# 标题..." | python build_docx.py --out out.docx          # 从 stdin 读

支持的 Markdown 子集（覆盖简历/求职信/讲稿所需）:
    # H1            文档主标题（姓名 / 求职信抬头 / 讲稿项目名）
    ## H2           章节标题（带浅色下边框）
    ### H3          小节标题
    普通段落         正文（求职信、讲稿口述内容）
    - / * 项         无序列表（简历 bullet）
    1. 项            有序列表
    **粗体** *斜体*  行内强调
    `code`          行内等宽（去掉反引号，保留文字）
    ---             分隔线（章节之间留白）
    > 引用           引用段（讲稿的"可能追问"等）

退出码: 0 成功; 2 入参错误; 3 渲染失败。
"""
from __future__ import annotations

import argparse
import re
import subprocess
import sys

# Windows 控制台默认 GBK，会让中文/emoji 输出乱码甚至抛 UnicodeEncodeError。
# 统一把 stdout/stderr 重配为 UTF-8，保证脚本在任何平台输出都正确。
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass


def _ensure_docx():
    """python-docx 缺失时自动安装，让脚本在裸环境也能跑通。"""
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        print("[build_docx] 未检测到 python-docx，正在自动安装…", file=sys.stderr)
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet", "python-docx"],
            check=True,
        )


_ensure_docx()

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

# 三种文档风格：仅调整正文行距与段后距，标题样式共用一套，保持视觉一致。
STYLES = {
    "resume": {"body_size": 10.5, "space_after": 3, "line_spacing": 1.08},
    "letter": {"body_size": 11.0, "space_after": 8, "line_spacing": 1.25},
    "pitch": {"body_size": 11.5, "space_after": 6, "line_spacing": 1.3},
}
ACCENT = RGBColor(0x1F, 0x49, 0x7D)  # 标题深蓝
RULE_COLOR = "BFBFBF"  # 章节下边框浅灰
CJK_FONT = "Microsoft YaHei"  # 东亚字体：Calibri 无中文字形，必须显式指定 w:eastAsia 否则渲染成豆腐块

_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _set_cjk(run, name: str = CJK_FONT):
    """给 run 设置 w:eastAsia 东亚字体。python-docx 设 font.name 只写 ascii/hAnsi，
    不写 eastAsia，导致中文回退到无字形的默认字体（Word 里显示为□）。"""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def _add_runs(paragraph, text: str, base_size: float):
    """把一行文本按 **粗体** / *斜体* / `code` 切成多个 run 写入段落。"""
    for token in _INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token)
        run.font.size = Pt(base_size)
        _set_cjk(run)


def _bottom_border(paragraph, color: str):
    """给段落加一条下边框，用作章节标题的分隔线。"""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def render(md: str, style: str) -> "Document":
    cfg = STYLES.get(style, STYLES["resume"])
    doc = Document()

    # 收窄页边距，简历/求职信单页更易排布。
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(50)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    # Normal 样式也写入 eastAsia，覆盖未显式设字体的继承 run（如标题继承）
    _n_rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    _n_rfonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(cfg["body_size"])
    normal.paragraph_format.line_spacing = cfg["line_spacing"]
    normal.paragraph_format.space_after = Pt(cfg["space_after"])

    lines = md.replace("\r\n", "\n").split("\n")
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        # 分隔线 → 一个带下边框的空段，制造章节间留白
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(2)
            _bottom_border(sep, RULE_COLOR)
            continue

        # 标题
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = ACCENT
                p.paragraph_format.space_after = Pt(4)
            elif level == 2:
                run = p.add_run(text.upper() if text.isascii() else text)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = ACCENT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                _bottom_border(p, RULE_COLOR)
            else:
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(cfg["body_size"] + 1)
                p.paragraph_format.space_before = Pt(4)
            _set_cjk(run)
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            _add_runs(p, stripped[1:].strip(), cfg["body_size"])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # 无序列表
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1), cfg["body_size"])
            p.paragraph_format.space_after = Pt(cfg["space_after"])
            continue

        # 有序列表
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1), cfg["body_size"])
            p.paragraph_format.space_after = Pt(cfg["space_after"])
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs(p, stripped, cfg["body_size"])

    return doc


def main() -> int:
    ap = argparse.ArgumentParser(description="Markdown → .docx (简历/求职信/讲稿)")
    ap.add_argument("--in", dest="src", help="输入 Markdown 文件；省略则读 stdin")
    ap.add_argument("--out", required=True, help="输出 .docx 路径")
    ap.add_argument(
        "--style",
        default="resume",
        choices=sorted(STYLES),
        help="排版风格：resume(紧凑) / letter(求职信) / pitch(讲稿)",
    )
    args = ap.parse_args()

    if args.src:
        try:
            with open(args.src, "r", encoding="utf-8") as fh:
                md = fh.read()
        except OSError as exc:
            print(f"[build_docx] 读不到输入文件: {exc}", file=sys.stderr)
            return 2
    else:
        md = sys.stdin.read()

    if not md.strip():
        print("[build_docx] 输入为空，没有内容可渲染。", file=sys.stderr)
        return 2

    try:
        doc = render(md, args.style)
        doc.save(args.out)
    except Exception as exc:  # noqa: BLE001
        print(f"[build_docx] 渲染失败: {exc}", file=sys.stderr)
        return 3

    # 简单自检：确认文件确实生成且非空
    import os

    size = os.path.getsize(args.out)
    print(f"[build_docx] 已生成 {args.out} ({size} bytes, style={args.style})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
